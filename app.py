import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time

# -----------------------------
# Config / API
# -----------------------------
GEMINI_API_KEY = "AIzaSyB1ZfX-LSaWQ1kEqpP8LCdcRZj6ym-Yty8"
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")

# -----------------------------
# Page Config & Custom CSS
# -----------------------------
st.set_page_config(
    page_title="Biochar PDD Generator",
    page_icon="🌿",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    /* Base font */
    html, body, [class*="css"] {
        font-family: 'Segoe UI', sans-serif;
    }

    /* Main background */
    .stApp {
        background-color: #f4f6f0;
    }

    /* Top header banner */
    .main-header {
        background: linear-gradient(135deg, #1a3c2e, #2d6a4f);
        color: white;
        padding: 2rem 2.5rem;
        border-radius: 12px;
        margin-bottom: 2rem;
    }
    .main-header h1 {
        font-size: 2rem;
        font-weight: 700;
        margin: 0;
    }
    .main-header p {
        font-size: 1rem;
        margin: 0.3rem 0 0 0;
        opacity: 0.85;
    }

    /* Step cards */
    .step-card {
        background: white;
        border-radius: 10px;
        padding: 1.5rem 2rem;
        margin-bottom: 1.5rem;
        border-left: 5px solid #2d6a4f;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }
    .step-title {
        font-size: 1.1rem;
        font-weight: 700;
        color: #1a3c2e;
        margin-bottom: 0.2rem;
    }
    .step-subtitle {
        font-size: 0.85rem;
        color: #6c757d;
        margin-bottom: 1rem;
    }

    /* Progress bar custom */
    .progress-wrapper {
        background: white;
        border-radius: 10px;
        padding: 1rem 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }

    /* Generate button */
    .stButton > button {
        background: linear-gradient(135deg, #1a3c2e, #2d6a4f);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 2.5rem;
        font-size: 1.05rem;
        font-weight: 600;
        width: 100%;
        transition: opacity 0.2s;
    }
    .stButton > button:hover {
        opacity: 0.9;
        color: white;
    }

    /* Section generation status */
    .section-status {
        background: #e8f5e9;
        border-left: 4px solid #2d6a4f;
        padding: 0.5rem 1rem;
        border-radius: 4px;
        margin: 0.3rem 0;
        font-size: 0.9rem;
        color: #1a3c2e;
    }

    /* PDD preview box */
    .pdd-preview {
        background: white;
        border-radius: 10px;
        padding: 2rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        line-height: 1.8;
    }

    /* Input labels */
    label {
        font-weight: 600 !important;
        color: #2c3e50 !important;
    }

    /* Selectbox and input */
    .stSelectbox > div, .stNumberInput > div, .stTextInput > div {
        border-radius: 6px !important;
    }

    /* Footer */
    .footer {
        text-align: center;
        padding: 1.5rem;
        color: #6c757d;
        font-size: 0.85rem;
        margin-top: 2rem;
    }

    /* Badge */
    .badge {
        background: #e8f5e9;
        color: #1a3c2e;
        padding: 0.2rem 0.7rem;
        border-radius: 20px;
        font-size: 0.78rem;
        font-weight: 600;
        display: inline-block;
        margin-left: 0.5rem;
    }

    /* Divider */
    hr {
        border: none;
        border-top: 1px solid #e0e0e0;
        margin: 1.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# -----------------------------
# Master Context (used in all prompts)
# -----------------------------
MASTER_CONTEXT = """
You are generating a Project Design Document (PDD) section for a biochar-based carbon removal project.

The document must align with the Puro.earth carbon removal framework for biochar.

Requirements:
- Maintain formal, technical, audit-ready language
- Focus on clarity and accuracy
- Avoid marketing language or vague statements
- Use ONLY the provided input data
- Do NOT fabricate or invent numerical values
- If any data is missing or incomplete, explicitly state the assumption
- Write in structured paragraphs with clear subheadings where appropriate
- Do not include the section title in the output
- Keep the section concise but sufficiently detailed (typically 150–300 words)
"""

# -----------------------------
# Section-wise Prompt Builders
# -----------------------------
def build_input_block(data):
    """Build a clean structured input block to inject into all prompts."""
    return f"""
Input Data:
- Project Name: {data.get('project_name', 'N/A')}
- Developer: {data.get('developer', 'N/A')}
- Location: {data.get('location', 'N/A')}
- Start Date: {data.get('start_date', 'N/A')}
- Crediting Period: {data.get('crediting_period', 'N/A')} years
- Registry / Standard: {data.get('registry', 'Puro.earth')}
- Annual Feedstock Input: {data.get('feedstock_input', 'N/A')} tons/year
- Annual Biochar Production: {data.get('biochar_production', 'N/A')} tons/year
- Operational Days per Year: {data.get('operational_days', 'N/A')} days
- Feedstock Type: {data.get('feedstock_type', 'N/A')}
- Feedstock Category: {data.get('feedstock_category', 'N/A')}
- Feedstock Source: {data.get('feedstock_source', 'N/A')}
- Competing Uses: {data.get('competing_uses', 'N/A')}
- Moisture Content: {data.get('moisture_content', 'N/A')}%
- Technology Type: {data.get('technology_type', 'N/A')}
- Reactor Type: {data.get('reactor_type', 'N/A')}
- Operating Temperature: {data.get('temperature', 'N/A')} °C
- Residence Time: {data.get('residence_time', 'N/A')} minutes
- Energy Source: {data.get('energy_source', 'N/A')}
- Biochar Yield: {data.get('biochar_yield', 'N/A')}%
- Fixed Carbon Content: {data.get('carbon_content', 'N/A')}%
- H/Corg Ratio: {data.get('hcorg_ratio', 'N/A')}
- Expected Stability: {data.get('stability', 'N/A')}
- End Use of Biochar: {data.get('end_use', 'N/A')}
- Baseline Scenario: {data.get('baseline_scenario', 'N/A')}
- Baseline Justification: {data.get('baseline_justification', 'N/A')}
- Transport Distance: {data.get('transport_distance', 'N/A')} km
- Transport Fuel Type: {data.get('transport_fuel', 'N/A')}
- Fossil Energy Used in Process: {data.get('fossil_energy', 'N/A')}
- Leakage Risk: {data.get('leakage_risk', 'N/A')}
- Financial Viability Without Carbon Revenue: {data.get('financial_viability', 'N/A')}
- Main Barrier: {data.get('main_barrier', 'N/A')}
- Monitoring Frequency: {data.get('monitoring_frequency', 'N/A')}
- Parameters Monitored: {', '.join(data.get('parameters_monitored', []))}
- Data Recording Method: {data.get('data_recording', 'N/A')}
"""

def prompt_project_description(data):
    return f"""{MASTER_CONTEXT}

Write the "Project Description" section of the PDD.

Include:
- Overview of the project activity and its purpose
- Location, scale, and operational context
- Feedstock used and its origin
- Technology employed for biochar production
- How the project achieves carbon removal

{build_input_block(data)}

Ensure:
- Clear explanation of the biochar production process
- Mention of carbon sequestration through stable biochar application
- Reference to Puro.earth framework alignment
- Formal tone throughout

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_feedstock_sustainability(data):
    return f"""{MASTER_CONTEXT}

Write the "Feedstock Sustainability" section of the PDD.

Include:
- Description of the feedstock and its classification
- Evidence that feedstock is waste or residue (no deforestation risk)
- Feedstock availability and supply chain
- Competing uses assessment and how they are managed
- Sustainability safeguards in place

{build_input_block(data)}

Ensure:
- Clear argument for feedstock eligibility under Puro.earth criteria
- Address any competing use concerns honestly
- Avoid generic statements — be specific to the feedstock type provided

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_technology_pyrolysis(data):
    return f"""{MASTER_CONTEXT}

Write the "Technology and Pyrolysis Process" section of the PDD.

Include:
- Description of the pyrolysis technology used
- Operating conditions (temperature, residence time, reactor type)
- Energy source and its implications for process emissions
- Process controls and quality assurance measures
- How the technology ensures consistent carbon stability

{build_input_block(data)}

Ensure:
- Technical accuracy in describing pyrolysis process
- Link between operating conditions and carbon permanence
- Mention of energy source and any fossil fuel implications
- Engineering credibility throughout

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_biochar_characteristics(data):
    return f"""{MASTER_CONTEXT}

Write the "Biochar Characteristics and Carbon Stability" section of the PDD.

Include:
- Physical and chemical properties of the biochar produced
- Carbon content and biochar yield
- H/Corg ratio as a stability indicator (lower H/Corg = higher stability)
- Expected permanence and carbon storage duration
- End use application and its suitability

{build_input_block(data)}

Ensure:
- H/Corg ratio is discussed as a permanence indicator (do not hardcode pass/fail thresholds)
- Carbon removal claim is clearly justified
- End use application is linked to long-term carbon storage
- Puro.earth stability criteria are referenced

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_baseline_scenario(data):
    return f"""{MASTER_CONTEXT}

Write the "Baseline Scenario" section of the PDD.

Include:
- Description of what would happen to the feedstock without this project
- Greenhouse gas emissions associated with the baseline scenario
- Why this baseline is realistic and representative
- Quantitative or qualitative emissions argument

{build_input_block(data)}

Ensure:
- Baseline matches the selected scenario exactly
- GHG emissions are explained logically (CO2, CH4, N2O where relevant)
- Justification is grounded in the provided reasoning
- Avoid generic or vague statements

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_emissions_leakage(data):
    return f"""{MASTER_CONTEXT}

Write the "Emissions and Leakage Assessment" section of the PDD.

Include:
- Project activity emissions (transport, fossil energy use)
- Leakage risks and how they are managed
- Net emissions calculation approach
- Any fossil fuel inputs and their GHG implications

{build_input_block(data)}

Ensure:
- All emission sources are identified and addressed
- Leakage risk level is explained with reasoning
- Transport emissions are quantified qualitatively based on distance and fuel type
- Honest assessment — do not minimize risks without justification

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_additionality(data):
    return f"""{MASTER_CONTEXT}

Write the "Additionality" section of the PDD.

Include:
- Assessment of financial viability without carbon revenue
- Identification of key barriers (financial, technological, market)
- Role of carbon finance in enabling the project
- Argument for why the project would not occur under business-as-usual

{build_input_block(data)}

Ensure:
- Additionality argument is logical and structured
- Barriers are specific — avoid exaggerated or unsubstantiated claims
- Carbon revenue is clearly positioned as the enabling factor
- Aligns with standard additionality reasoning used in carbon markets

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

def prompt_monitoring_plan(data):
    return f"""{MASTER_CONTEXT}

Write the "Monitoring Plan" section of the PDD.

Include:
- Parameters to be monitored and why each is important
- Monitoring frequency and schedule
- Data collection and recording methods
- Quality assurance and verification procedures
- Roles and responsibilities for monitoring

{build_input_block(data)}

Ensure:
- All monitored parameters are clearly explained
- Monitoring approach supports verifiability
- Data recording method is described with enough detail
- Aligned with Puro.earth verification requirements

Output format:
- Use clear subheadings where appropriate
- Write in structured paragraphs (avoid excessive bullet points)
- Do not include the section title in the output
- If any required information is not provided, explicitly state the assumption instead of inventing data
"""

# -----------------------------
# Gemini API Call
# -----------------------------
def call_gemini(prompt, retries=3):
    """Call Gemini API with retry logic."""
    for attempt in range(retries):
        try:
            response = model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            if attempt < retries - 1:
                time.sleep(2)
            else:
                return f"[Error generating this section: {str(e)}]"

# -----------------------------
# PDD Assembly
# -----------------------------
def assemble_pdd(sections, data):
    """Assemble all sections into a full PDD string."""
    pdd = f"""
PROJECT DESIGN DOCUMENT (PDD)
Biochar Carbon Removal Project | Puro.earth Framework

Project: {data.get('project_name', 'N/A')}
Developer: {data.get('developer', 'N/A')}
Location: {data.get('location', 'N/A')}
Standard: {data.get('registry', 'Puro.earth')}
Crediting Period: {data.get('crediting_period', 'N/A')} years
Date: {data.get('start_date', 'N/A')}

{'='*60}

1. PROJECT DESCRIPTION
{'-'*40}
{sections.get('project_description', '')}

2. FEEDSTOCK SUSTAINABILITY
{'-'*40}
{sections.get('feedstock_sustainability', '')}

3. TECHNOLOGY AND PYROLYSIS PROCESS
{'-'*40}
{sections.get('technology_pyrolysis', '')}

4. BIOCHAR CHARACTERISTICS AND CARBON STABILITY
{'-'*40}
{sections.get('biochar_characteristics', '')}

5. BASELINE SCENARIO
{'-'*40}
{sections.get('baseline_scenario', '')}

6. EMISSIONS AND LEAKAGE ASSESSMENT
{'-'*40}
{sections.get('emissions_leakage', '')}

7. ADDITIONALITY
{'-'*40}
{sections.get('additionality', '')}

8. MONITORING PLAN
{'-'*40}
{sections.get('monitoring_plan', '')}

{'='*60}
Generated using Biochar PDD Generator | Puro.earth Aligned
"""
    return pdd

# -----------------------------
# DOCX Export
# -----------------------------
def generate_docx(sections, data):
    """Generate a professional Word document from PDD sections."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading('Project Design Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x2e)
    title_run.font.size = Pt(22)

    subtitle = doc.add_paragraph('Biochar Carbon Removal Project | Puro.earth Framework')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True

    doc.add_paragraph('')

    # Project info table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    info_rows = [
        ('Project Name', data.get('project_name', 'N/A')),
        ('Developer', data.get('developer', 'N/A')),
        ('Location', data.get('location', 'N/A')),
        ('Standard', data.get('registry', 'Puro.earth')),
        ('Crediting Period', f"{data.get('crediting_period', 'N/A')} years"),
        ('Start Date', str(data.get('start_date', 'N/A'))),
    ]
    for i, (key, value) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1a, 0x3c, 0x2e)

    doc.add_paragraph('')

    # Sections
    section_titles = [
        ('project_description', '1. Project Description'),
        ('feedstock_sustainability', '2. Feedstock Sustainability'),
        ('technology_pyrolysis', '3. Technology and Pyrolysis Process'),
        ('biochar_characteristics', '4. Biochar Characteristics and Carbon Stability'),
        ('baseline_scenario', '5. Baseline Scenario'),
        ('emissions_leakage', '6. Emissions and Leakage Assessment'),
        ('additionality', '7. Additionality'),
        ('monitoring_plan', '8. Monitoring Plan'),
    ]

    for key, title_text in section_titles:
        # Section heading
        heading = doc.add_heading(title_text, level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x2e)
        heading_run.font.size = Pt(14)

        # Section content
        content = sections.get(key, '')
        for para in content.split('\n'):
            para = para.strip()
            if not para:
                continue
            if para.startswith('**') and para.endswith('**'):
                # Bold subheading
                p = doc.add_paragraph()
                run = p.add_run(para.replace('**', ''))
                run.bold = True
                run.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
            else:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph('')

    # Footer
    footer_para = doc.add_paragraph('Generated by Biochar Solutions | Puro.earth Aligned | For review purposes only')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
    footer_run.font.italic = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------
# Streamlit UI
# -----------------------------

# Header
st.markdown("""
<div class="main-header">
    <h1>🌿 Biochar PDD Generator</h1>
    <p>Puro.earth Framework &nbsp;|&nbsp; AI-Powered Project Design Document &nbsp;|&nbsp; Audit-Ready Output</p>
</div>
""", unsafe_allow_html=True)

# Intro
st.markdown("""
Fill in the form below across all 9 sections. Once complete, click **Generate PDD** to produce a 
full, Puro.earth-aligned Project Design Document. The tool uses section-wise AI generation to ensure 
consistency and methodology compliance.
""")

st.markdown("---")

# -----------------------------
# SECTION 1: Project Overview
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">📋 Section 1 — Project Overview</div>
    <div class="step-subtitle">Basic identity and context of the project</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    project_name = st.text_input("Project Name", placeholder="e.g. GreenChar Rajasthan")
    location = st.text_input("Location (State, Country)", placeholder="e.g. Rajasthan, India")
    crediting_period = st.number_input("Crediting Period (years)", min_value=1, max_value=30, value=10)
with col2:
    developer = st.text_input("Project Developer / Company", placeholder="e.g. Biochar Solutions India Pvt Ltd")
    start_date = st.date_input("Project Start Date")
    registry = st.selectbox("Registry / Standard", ["Puro.earth"])

st.markdown("---")

# -----------------------------
# SECTION 2: Project Scale
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">📏 Section 2 — Project Scale</div>
    <div class="step-subtitle">Annual throughput and operational capacity</div>
</div>
""", unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1:
    feedstock_input = st.number_input("Annual Feedstock Input (tons/year)", min_value=1, value=5000)
with col2:
    biochar_production = st.number_input("Annual Biochar Production (tons/year)", min_value=1, value=1500)
with col3:
    operational_days = st.number_input("Operational Days per Year", min_value=1, max_value=365, value=300)

st.markdown("---")

# -----------------------------
# SECTION 3: Feedstock Details
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">🌾 Section 3 — Feedstock Details</div>
    <div class="step-subtitle">Critical for sustainability assessment and baseline scenario</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    feedstock_type = st.text_input("Feedstock Type", placeholder="e.g. Rice husk, Wheat straw, Wood chips")
    feedstock_source = st.selectbox("Feedstock Source", [
        "Own operations",
        "Third-party suppliers",
        "Mixed"
    ])
    moisture_content = st.number_input("Moisture Content (%)", min_value=0, max_value=100, value=12)
with col2:
    feedstock_category = st.selectbox("Feedstock Category", [
        "Agricultural residue",
        "Forestry residue",
        "Industrial biomass waste"
    ])
    competing_uses = st.selectbox("Competing Uses of Feedstock", [
        "No significant competing use",
        "Used as fuel",
        "Used as fodder",
        "Other"
    ])

st.markdown("---")

# -----------------------------
# SECTION 4: Pyrolysis Technology
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">🔥 Section 4 — Pyrolysis Technology</div>
    <div class="step-subtitle">Defines carbon stability and process emissions</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    technology_type = st.selectbox("Technology Type", [
        "Slow pyrolysis",
        "Fast pyrolysis"
    ])
    temperature = st.number_input("Operating Temperature (°C)", min_value=200, max_value=1200, value=500)
    energy_source = st.selectbox("Energy Source", [
        "Renewable",
        "Fossil",
        "Mixed"
    ])
with col2:
    reactor_type = st.text_input("Reactor Type", placeholder="e.g. Rotary kiln, Fixed bed, TLUD")
    residence_time = st.number_input("Residence Time (minutes)", min_value=1, value=60)

st.markdown("---")

# -----------------------------
# SECTION 5: Biochar Output
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">⚗️ Section 5 — Biochar Characteristics</div>
    <div class="step-subtitle">Core of the carbon removal claim — stability and permanence</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    biochar_yield = st.number_input("Biochar Yield (%)", min_value=1, max_value=100, value=30)
    hcorg_ratio = st.number_input("H/Corg Ratio (molar)", min_value=0.0, max_value=1.0, step=0.01, value=0.38,
                                   help="Lower H/Corg indicates higher carbon stability. Used as a permanence indicator.")
    stability = st.selectbox("Expected Carbon Stability", [
        "High (>100 years)",
        "Medium (50–100 years)",
        "Unknown"
    ])
with col2:
    carbon_content = st.number_input("Fixed Carbon Content (%)", min_value=1, max_value=100, value=72)
    end_use = st.selectbox("End Use of Biochar", [
        "Soil application",
        "Construction materials",
        "Other"
    ])

st.markdown("---")

# -----------------------------
# SECTION 6: Baseline Scenario
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">📉 Section 6 — Baseline Scenario</div>
    <div class="step-subtitle">What happens to the feedstock without this project?</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    baseline_scenario = st.selectbox("Baseline Scenario", [
        "Open burning",
        "Landfill disposal",
        "Natural decomposition"
    ])
with col2:
    baseline_justification = st.selectbox("Baseline Justification", [
        "Common practice in the region",
        "Lack of alternative waste management infrastructure",
        "Economically most viable option",
        "Regulatory absence",
        "Other (see project description)"
    ])

st.markdown("---")

# -----------------------------
# SECTION 7: Emissions & Leakage
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">💨 Section 7 — Emissions & Leakage</div>
    <div class="step-subtitle">Transport, fossil inputs and leakage risk assessment</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    transport_distance = st.number_input("Average Transport Distance (km)", min_value=0, value=50)
    fossil_energy = st.selectbox("Fossil Energy Used in Pyrolysis Process?", ["No", "Yes"])
with col2:
    transport_fuel = st.selectbox("Transport Fuel Type", ["Diesel", "Petrol", "Electric"])
    leakage_risk = st.selectbox("Leakage Risk", ["Low", "Medium", "High"])

st.markdown("---")

# -----------------------------
# SECTION 8: Additionality
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">✅ Section 8 — Additionality</div>
    <div class="step-subtitle">Why this project would not happen without carbon finance</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    financial_viability = st.selectbox("Financial Viability Without Carbon Revenue", [
        "Not viable",
        "Marginally viable",
        "Fully viable"
    ])
with col2:
    main_barrier = st.selectbox("Main Barrier to Implementation", [
        "Financial",
        "Technological",
        "Market"
    ])

st.markdown("---")

# -----------------------------
# SECTION 9: Monitoring Plan
# -----------------------------
st.markdown("""
<div class="step-card">
    <div class="step-title">📊 Section 9 — Monitoring Plan</div>
    <div class="step-subtitle">How carbon removals will be measured and verified</div>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    monitoring_frequency = st.selectbox("Monitoring Frequency", [
        "Continuous",
        "Daily",
        "Monthly"
    ])
    data_recording = st.selectbox("Data Recording Method", [
        "Manual logs",
        "Digital system",
        "Automated sensors"
    ])
with col2:
    parameters_monitored = st.multiselect(
        "Parameters Monitored",
        ["Feedstock quantity", "Biochar production", "Carbon content", "Temperature"],
        default=["Feedstock quantity", "Biochar production", "Carbon content", "Temperature"]
    )

st.markdown("---")

# -----------------------------
# Generate PDD Button
# -----------------------------
st.markdown("### 🚀 Generate PDD")
st.markdown("All sections will be generated separately using AI, then assembled into a full document.")

if st.button("⚡ Generate Full PDD"):

    # Validate required fields
    if not project_name or not developer or not location:
        st.error("Please fill in at least Project Name, Developer, and Location before generating.")
    else:
        # Collect all inputs
        user_data = {
            "project_name": project_name,
            "developer": developer,
            "location": location,
            "start_date": str(start_date),
            "crediting_period": crediting_period,
            "registry": registry,
            "feedstock_input": feedstock_input,
            "biochar_production": biochar_production,
            "operational_days": operational_days,
            "feedstock_type": feedstock_type,
            "feedstock_category": feedstock_category,
            "feedstock_source": feedstock_source,
            "competing_uses": competing_uses,
            "moisture_content": moisture_content,
            "technology_type": technology_type,
            "reactor_type": reactor_type,
            "temperature": temperature,
            "residence_time": residence_time,
            "energy_source": energy_source,
            "biochar_yield": biochar_yield,
            "carbon_content": carbon_content,
            "hcorg_ratio": hcorg_ratio,
            "stability": stability,
            "end_use": end_use,
            "baseline_scenario": baseline_scenario,
            "baseline_justification": baseline_justification,
            "transport_distance": transport_distance,
            "transport_fuel": transport_fuel,
            "fossil_energy": fossil_energy,
            "leakage_risk": leakage_risk,
            "financial_viability": financial_viability,
            "main_barrier": main_barrier,
            "monitoring_frequency": monitoring_frequency,
            "parameters_monitored": parameters_monitored,
            "data_recording": data_recording,
        }

        # Section-wise generation with progress
        sections = {}
        section_list = [
            ("project_description",     "1. Project Description",                    prompt_project_description),
            ("feedstock_sustainability", "2. Feedstock Sustainability",               prompt_feedstock_sustainability),
            ("technology_pyrolysis",     "3. Technology and Pyrolysis Process",       prompt_technology_pyrolysis),
            ("biochar_characteristics",  "4. Biochar Characteristics & Stability",    prompt_biochar_characteristics),
            ("baseline_scenario",        "5. Baseline Scenario",                      prompt_baseline_scenario),
            ("emissions_leakage",        "6. Emissions and Leakage Assessment",       prompt_emissions_leakage),
            ("additionality",            "7. Additionality",                          prompt_additionality),
            ("monitoring_plan",          "8. Monitoring Plan",                        prompt_monitoring_plan),
        ]

        st.markdown("**Generating sections...**")
        progress_bar = st.progress(0)
        status_placeholder = st.empty()

        for i, (key, title, prompt_fn) in enumerate(section_list):
            status_placeholder.markdown(f"""
            <div class="section-status">⏳ Generating: {title}</div>
            """, unsafe_allow_html=True)
            sections[key] = call_gemini(prompt_fn(user_data))
            progress_bar.progress((i + 1) / len(section_list))
            time.sleep(0.5)  # brief pause between API calls

        status_placeholder.markdown("""
        <div class="section-status">✅ All sections generated successfully!</div>
        """, unsafe_allow_html=True)

        # Assemble PDD
        full_pdd = assemble_pdd(sections, user_data)

        st.markdown("---")
        st.markdown("### 📄 Generated PDD Preview")

        # Preview
        with st.expander("👁️ Click to preview full PDD", expanded=True):
            st.markdown(f"""<div class="pdd-preview">{full_pdd.replace(chr(10), '<br>')}</div>""",
                        unsafe_allow_html=True)

        # Download buttons
        st.markdown("### ⬇️ Download")
        col1, col2 = st.columns(2)

        with col1:
            # Plain text download
            st.download_button(
                label="📥 Download as .txt",
                data=full_pdd,
                file_name=f"{project_name.replace(' ', '_')}_PDD.txt",
                mime="text/plain"
            )

        with col2:
            # Word doc download
            docx_buffer = generate_docx(sections, user_data)
            st.download_button(
                label="📄 Download as Word (.docx)",
                data=docx_buffer,
                file_name=f"{project_name.replace(' ', '_')}_PDD.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# -----------------------------
# Footer
# -----------------------------
st.markdown("---")
st.markdown("""
<div class="footer">
    🌿 Biochar PDD Generator &nbsp;|&nbsp; Puro.earth Framework &nbsp;|&nbsp; 
    Built for carbon project developers &nbsp;|&nbsp; 
    <strong>Made with ❤️ by Mayank Kumar Sharma</strong>
</div>
""", unsafe_allow_html=True)
